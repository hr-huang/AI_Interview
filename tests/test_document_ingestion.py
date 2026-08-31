import io
import unittest

import pymupdf
from docx import Document

from profile_agent.web.document_ingestion import (
    MAX_FILE_BYTES,
    DocumentExtractor,
    OcrLine,
)


class FakeOcr:
    def __init__(self) -> None:
        self.calls = 0

    def recognize(self, image_bytes: bytes) -> list[OcrLine]:
        self.calls += 1
        self.last_image = image_bytes
        return [
            OcrLine(
                text=(
                    "扫描页中的 LangGraph 项目：负责 Agent Workflow 设计、"
                    "checkpoint 恢复与上线监控。"
                ),
                confidence=0.98,
                y=10,
            )
        ]


class DocumentIngestionTest(unittest.TestCase):
    def test_text_pdf_does_not_use_ocr(self) -> None:
        pdf = pymupdf.open()
        page = pdf.new_page()
        page.insert_text(
            (50, 50),
            "AI Agent Workflow LangGraph FastAPI project experience and delivery",
        )
        content = pdf.tobytes()
        ocr = FakeOcr()

        result = DocumentExtractor(ocr_engine=ocr).extract("resume.pdf", content)

        self.assertIn("LangGraph", result.text)
        self.assertEqual(result.file_type, "pdf")
        self.assertEqual(result.used_ocr_pages, [])
        self.assertEqual(ocr.calls, 0)

    def test_blank_pdf_uses_ocr_for_only_blank_page(self) -> None:
        pdf = pymupdf.open()
        pdf.new_page()
        ocr = FakeOcr()

        result = DocumentExtractor(ocr_engine=ocr).extract(
            "scan.pdf",
            pdf.tobytes(),
        )

        self.assertIn("扫描页", result.text)
        self.assertEqual(result.used_ocr_pages, [1])
        self.assertEqual(ocr.calls, 1)
        self.assertTrue(ocr.last_image.startswith(b"\x89PNG"))

    def test_docx_reads_paragraphs_and_tables(self) -> None:
        stream = io.BytesIO()
        document = Document()
        document.add_paragraph(
            "候选人简介：负责 AI Agent 产品的后端开发与工作流落地。"
        )
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Agent 项目：LangGraph 状态管理与失败恢复。"
        document.save(stream)

        result = DocumentExtractor(ocr_engine=FakeOcr()).extract(
            "resume.docx",
            stream.getvalue(),
        )

        self.assertIn("候选人简介", result.text)
        self.assertIn("Agent 项目", result.text)
        self.assertEqual(result.file_type, "docx")
        self.assertEqual(result.used_ocr_pages, [])

    def test_txt_accepts_utf8_and_gb18030(self) -> None:
        text = "AI 应用工程师，具备 LangGraph 工作流、RAG 评估和可观测性落地经验。"
        extractor = DocumentExtractor(ocr_engine=FakeOcr())

        utf8 = extractor.extract("resume.txt", text.encode("utf-8"))
        gb18030 = extractor.extract("resume.txt", text.encode("gb18030"))

        self.assertEqual(utf8.text, text)
        self.assertEqual(gb18030.text, text)
        self.assertEqual(utf8.file_type, "txt")

    def test_old_doc_is_rejected(self) -> None:
        with self.assertRaisesRegex(ValueError, "DOCX 或 PDF"):
            DocumentExtractor(ocr_engine=FakeOcr()).extract("resume.doc", b"x")

    def test_content_signature_wins_over_misleading_suffix(self) -> None:
        pdf = pymupdf.open()
        page = pdf.new_page()
        page.insert_text(
            (50, 50),
            "LangGraph workflow evaluation reliability monitoring experience",
        )
        result = DocumentExtractor(ocr_engine=FakeOcr()).extract(
            "resume.txt",
            pdf.tobytes(),
        )
        self.assertEqual(result.file_type, "pdf")

    def test_empty_oversized_and_unknown_files_are_rejected(self) -> None:
        extractor = DocumentExtractor(ocr_engine=FakeOcr())
        with self.assertRaisesRegex(ValueError, "空或超过"):
            extractor.extract("resume.txt", b"")
        with self.assertRaisesRegex(ValueError, "空或超过"):
            extractor.extract("resume.txt", b"x" * (MAX_FILE_BYTES + 1))
        with self.assertRaisesRegex(ValueError, "仅支持"):
            extractor.extract("resume.png", b"not-an-image")


if __name__ == "__main__":
    unittest.main()
