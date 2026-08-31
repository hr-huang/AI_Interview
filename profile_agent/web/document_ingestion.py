from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from typing import Protocol

import pymupdf
from docx import Document

MAX_FILE_BYTES = 5 * 1024 * 1024
MAX_PDF_PAGES = 10


@dataclass(frozen=True)
class OcrLine:
    text: str
    confidence: float
    y: float


class OcrEngine(Protocol):
    def recognize(self, image_bytes: bytes) -> list[OcrLine]: ...


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    file_type: str
    used_ocr_pages: list[int]


def _usable(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    replacement_ratio = text.count("�") / max(1, len(text))
    return len(compact) >= 30 and replacement_ratio < 0.05


class RapidOcrEngine:
    def __init__(self) -> None:
        from rapidocr import RapidOCR

        self._engine = RapidOCR()

    def recognize(self, image_bytes: bytes) -> list[OcrLine]:
        result = self._engine(image_bytes)
        boxes = [] if result.boxes is None else result.boxes
        texts = () if result.txts is None else result.txts
        scores = () if result.scores is None else result.scores
        lines = []
        for box, text, confidence in zip(boxes, texts, scores):
            y = min(point[1] for point in box)
            lines.append(
                OcrLine(
                    text=str(text),
                    confidence=float(confidence),
                    y=float(y),
                )
            )
        return sorted(lines, key=lambda line: line.y)


class DocumentExtractor:
    def __init__(self, ocr_engine: OcrEngine | None = None) -> None:
        self._ocr_engine = ocr_engine

    @property
    def ocr_engine(self) -> OcrEngine:
        if self._ocr_engine is None:
            self._ocr_engine = RapidOcrEngine()
        return self._ocr_engine

    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        if not content or len(content) > MAX_FILE_BYTES:
            raise ValueError("简历文件为空或超过 5 MiB")
        suffix = filename.lower().rsplit(".", 1)[-1]
        if suffix == "doc":
            raise ValueError("旧版 DOC 不支持，请转换为 DOCX 或 PDF")
        if content.startswith(b"%PDF-"):
            return self._extract_pdf(content)
        if content.startswith(b"PK") and self._is_docx(content):
            return self._extract_docx(content)
        if suffix == "txt":
            return self._extract_txt(content)
        raise ValueError("仅支持 PDF、DOCX 或 TXT 简历")

    @staticmethod
    def _is_docx(content: bytes) -> bool:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
        except zipfile.BadZipFile:
            return False
        return {
            "[Content_Types].xml",
            "word/document.xml",
        }.issubset(names)

    def _extract_pdf(self, content: bytes) -> ExtractedDocument:
        try:
            document = pymupdf.open(stream=content, filetype="pdf")
        except (pymupdf.FileDataError, RuntimeError) as error:
            raise ValueError("PDF 文件已损坏或格式无效") from error
        try:
            if document.page_count > MAX_PDF_PAGES:
                raise ValueError("PDF 最多支持 10 页")
            pages: list[str] = []
            used_ocr_pages: list[int] = []
            for page_number, page in enumerate(document, start=1):
                native_text = self._native_pdf_text(page)
                if _usable(native_text):
                    pages.append(native_text)
                    continue
                image = page.get_pixmap(dpi=180).tobytes("png")
                ocr_text = "\n".join(
                    line.text.strip()
                    for line in self.ocr_engine.recognize(image)
                    if line.confidence >= 0.5 and line.text.strip()
                )
                if not _usable(ocr_text):
                    raise ValueError(
                        f"PDF 第 {page_number} 页无法提取可用文本"
                    )
                pages.append(ocr_text)
                used_ocr_pages.append(page_number)
            return ExtractedDocument(
                text="\n\n".join(pages),
                file_type="pdf",
                used_ocr_pages=used_ocr_pages,
            )
        finally:
            document.close()

    @staticmethod
    def _native_pdf_text(page: pymupdf.Page) -> str:
        blocks = page.get_text("blocks", sort=True)
        return "\n".join(
            str(block[4]).strip()
            for block in blocks
            if len(block) > 6
            and block[6] == 0
            and str(block[4]).strip()
        )

    @staticmethod
    def _extract_docx(content: bytes) -> ExtractedDocument:
        try:
            document = Document(io.BytesIO(content))
        except (ValueError, zipfile.BadZipFile, KeyError) as error:
            raise ValueError("DOCX 文件已损坏或格式无效") from error
        parts = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        parts.extend(
            cell.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        )
        text = "\n".join(parts)
        if not _usable(text):
            raise ValueError("DOCX 未提取到可用文本")
        return ExtractedDocument(text=text, file_type="docx", used_ocr_pages=[])

    @staticmethod
    def _extract_txt(content: bytes) -> ExtractedDocument:
        for encoding in ("utf-8-sig", "gb18030"):
            try:
                text = content.decode(encoding)
            except UnicodeDecodeError:
                continue
            if _usable(text):
                return ExtractedDocument(
                    text=text.strip(),
                    file_type="txt",
                    used_ocr_pages=[],
                )
        raise ValueError("TXT 编码不支持或未提取到可用文本")
